#!/usr/bin/env python3
"""
Convert .docx files to Markdown format.

Usage:
    python3 docx_to_markdown.py input.docx [output.md]

If output.md is not specified, outputs to stdout.
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def get_heading_level(paragraph):
    """Determine if paragraph is a heading and what level."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""

    # Check for heading styles
    if 'heading' in style_name:
        # Extract number from "Heading 1", "Heading 2", etc.
        for i in range(1, 7):
            if f'heading {i}' in style_name or style_name == f'heading{i}':
                return i

    # Check for title style
    if 'title' in style_name:
        return 1

    # Check for subtitle
    if 'subtitle' in style_name:
        return 2

    return 0


def is_list_paragraph(paragraph):
    """Check if paragraph is a list item."""
    # Check style name
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if 'list' in style_name:
        return True

    # Check for numbering
    if paragraph._element.pPr is not None:
        numPr = paragraph._element.pPr.numPr
        if numPr is not None:
            return True

    return False


def get_list_level(paragraph):
    """Get the indentation level of a list item."""
    if paragraph._element.pPr is not None:
        numPr = paragraph._element.pPr.numPr
        if numPr is not None:
            ilvl = numPr.ilvl
            if ilvl is not None:
                return int(ilvl.val)
    return 0


def is_numbered_list(paragraph):
    """Check if paragraph is a numbered list item."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    return 'number' in style_name or 'ordered' in style_name


def format_run(run):
    """Format a run with appropriate markdown."""
    text = run.text
    if not text:
        return ""

    # Apply formatting
    if run.bold and run.italic:
        text = f"***{text}***"
    elif run.bold:
        text = f"**{text}**"
    elif run.italic:
        text = f"*{text}*"

    # Underline is not standard markdown, skip
    # Strikethrough
    if run.font.strike:
        text = f"~~{text}~~"

    return text


def paragraph_to_markdown(paragraph, list_counters=None):
    """Convert a paragraph to markdown."""
    if list_counters is None:
        list_counters = {}

    # Get text with formatting
    text_parts = []
    for run in paragraph.runs:
        text_parts.append(format_run(run))
    text = "".join(text_parts).strip()

    if not text:
        return ""

    # Check for heading
    heading_level = get_heading_level(paragraph)
    if heading_level > 0:
        return "#" * heading_level + " " + text

    # Check for list
    if is_list_paragraph(paragraph):
        level = get_list_level(paragraph)
        indent = "  " * level

        if is_numbered_list(paragraph):
            # Track numbered list counters per level
            if level not in list_counters:
                list_counters[level] = 0
            list_counters[level] += 1
            # Reset deeper levels
            for l in list(list_counters.keys()):
                if l > level:
                    del list_counters[l]
            return f"{indent}{list_counters[level]}. {text}"
        else:
            return f"{indent}- {text}"

    # Regular paragraph
    return text


def table_to_markdown(table):
    """Convert a table to markdown."""
    if not table.rows:
        return ""

    markdown_rows = []

    for i, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get cell text (may span multiple paragraphs)
            cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            # Escape pipe characters
            cell_text = cell_text.replace("|", "\\|")
            cells.append(cell_text)

        markdown_rows.append("| " + " | ".join(cells) + " |")

        # Add header separator after first row
        if i == 0:
            separator = "| " + " | ".join(["---"] * len(cells)) + " |"
            markdown_rows.append(separator)

    return "\n".join(markdown_rows)


def convert_docx_to_markdown(docx_path):
    """Convert a .docx file to markdown string."""
    doc = Document(docx_path)

    markdown_parts = []
    list_counters = {}
    prev_was_list = False

    for element in doc.element.body:
        # Handle paragraphs
        if element.tag.endswith('p'):
            # Find the corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    md = paragraph_to_markdown(para, list_counters)

                    # Track list state for proper spacing
                    is_list = is_list_paragraph(para)
                    if not is_list and prev_was_list:
                        # Reset list counters when exiting list
                        list_counters = {}
                    prev_was_list = is_list

                    if md:
                        markdown_parts.append(md)
                    elif markdown_parts and markdown_parts[-1] != "":
                        # Add blank line for empty paragraphs (but not multiple)
                        markdown_parts.append("")
                    break

        # Handle tables
        elif element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    md = table_to_markdown(table)
                    if md:
                        if markdown_parts and markdown_parts[-1] != "":
                            markdown_parts.append("")
                        markdown_parts.append(md)
                        markdown_parts.append("")
                    break

    # Clean up multiple blank lines
    result = []
    prev_blank = False
    for part in markdown_parts:
        if part == "":
            if not prev_blank:
                result.append(part)
            prev_blank = True
        else:
            result.append(part)
            prev_blank = False

    # Remove trailing blank lines
    while result and result[-1] == "":
        result.pop()

    return "\n\n".join(line if line else "" for line in result)


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 docx_to_markdown.py input.docx [output.md]", file=sys.stderr)
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if not input_path.suffix.lower() == '.docx':
        print(f"Error: File must be a .docx file: {input_path}", file=sys.stderr)
        sys.exit(1)

    try:
        markdown = convert_docx_to_markdown(input_path)
    except Exception as e:
        print(f"Error converting file: {e}", file=sys.stderr)
        sys.exit(1)

    # Output
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
        output_path.write_text(markdown, encoding='utf-8')
        print(f"Converted: {input_path} -> {output_path}", file=sys.stderr)
    else:
        print(markdown)


if __name__ == "__main__":
    main()
